"""Resume parsing service for PDF and DOCX files."""

import os
import re
import logging
from pathlib import Path
from typing import Optional, Tuple

import fitz  # PyMuPDF
from docx import Document

from app.models.candidate import ResumeData, Education, Experience
from app.services.usf_client import usf_client

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parser for extracting text and structured data from resumes."""
    
    SUPPORTED_FORMATS = [".pdf", ".docx", ".doc"]
    
    def __init__(self):
        self.usf_client = usf_client
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text content from a PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
        
        return self._clean_text(text)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            doc = Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
        
        return self._clean_text(text)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a resume file (PDF or DOCX).
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')
        # Normalize multiple newlines to double newline
        text = re.sub(r'\n\s*\n', '\n\n', text)
        # Remove excessive spaces (but preserve newlines and double spaces for name detection)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Preserve double spaces (used in spaced-out names) but clean triple+ spaces
            line = re.sub(r' {3,}', '  ', line)
            cleaned_lines.append(line.strip())
        return '\n'.join(cleaned_lines).strip()
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_pattern, text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        phone_patterns = [
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{1,3}[-.\s]?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{2,4}',
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL from text."""
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        match = re.search(linkedin_pattern, text.lower())
        if match:
            return f"https://www.{match.group(0)}"
        return None
    
    def _basic_skill_extraction(self, text: str) -> list:
        """Extract skills using keyword matching as fallback."""
        common_skills = [
            # Programming languages
            "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "go", 
            "rust", "php", "swift", "kotlin", "scala", "r", "matlab", "sql",
            # Frameworks
            "react", "angular", "vue", "node.js", "django", "flask", "spring", 
            "express", "fastapi", ".net", "rails", "laravel", "tensorflow", "pytorch",
            "keras", "scikit-learn", "pandas", "numpy", "spark", "hadoop",
            # Tools
            "git", "docker", "kubernetes", "aws", "azure", "gcp", "jenkins", 
            "terraform", "ansible", "linux", "mongodb", "postgresql", "mysql",
            "redis", "elasticsearch", "kafka", "rabbitmq", "jira", "confluence",
            # Skills
            "machine learning", "deep learning", "data science", "data analysis",
            "agile", "scrum", "devops", "ci/cd", "rest api", "graphql",
            "microservices", "cloud computing", "project management", "nlp",
            "natural language processing", "computer vision", "neural networks"
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return found_skills
    
    def _extract_experience_years(self, text: str) -> float:
        """Extract total years of experience from resume text."""
        total_years = 0.0
        text_lower = text.lower()
        
        # Pattern 1: "X years of experience" or "X+ years experience"
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
            r'(?:experience|exp)\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
            r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:in|as|of)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = float(match)
                    if years > total_years and years < 50:  # Sanity check
                        total_years = years
                except:
                    pass
        
        # Pattern 2: Date ranges like "2018 - 2023" or "Jan 2019 - Present"
        if total_years == 0:
            date_patterns = [
                r'(20\d{2})\s*[-–]\s*(20\d{2}|present|current)',
                r'(19\d{2})\s*[-–]\s*(20\d{2}|present|current)',
            ]
            
            years_set = set()
            for pattern in date_patterns:
                matches = re.findall(pattern, text_lower)
                for start, end in matches:
                    try:
                        start_year = int(start)
                        if 'present' in end or 'current' in end:
                            end_year = 2026
                        else:
                            end_year = int(end)
                        years_set.add((start_year, end_year))
                    except:
                        pass
            
            # Calculate total from date ranges (avoid double counting)
            if years_set:
                min_year = min(s for s, e in years_set)
                max_year = max(e for s, e in years_set)
                total_years = max_year - min_year
        
        return min(total_years, 40.0)  # Cap at 40 years
    
    def _extract_education(self, text: str) -> list:
        """Extract education information from resume text."""
        education_list = []
        text_lower = text.lower()
        
        # Look for education section
        edu_section = ""
        edu_markers = ['education', 'academic', 'qualification', 'degree']
        for marker in edu_markers:
            if marker in text_lower:
                idx = text_lower.find(marker)
                edu_section = text_lower[idx:idx+500]
                break
        
        search_text = edu_section if edu_section else text_lower
        
        # Degree patterns - more specific to avoid false matches
        degree_patterns = [
            (r"master(?:'s)?\s+(?:of\s+)?(?:science|arts|engineering|business|technology)\s+(?:in\s+)?([\w\s]{3,30})", "Master's"),
            (r"m\.?s\.?\s+(?:in\s+)?([\w\s]{3,30})", "Master's"),
            (r"mba", "MBA"),
            (r"bachelor(?:'s)?\s+(?:of\s+)?(?:science|arts|engineering|technology)\s+(?:in\s+)?([\w\s]{3,30})", "Bachelor's"),
            (r"b\.?s\.?\s+(?:in\s+)?([\w\s]{3,30})", "Bachelor's"),
            (r"b\.?tech\s+(?:in\s+)?([\w\s]{3,30})", "Bachelor's"),
            (r"ph\.?d\.?\s+(?:in\s+)?([\w\s]{3,30})", "PhD"),
            (r"doctorate\s+(?:in\s+)?([\w\s]{3,30})", "PhD"),
        ]
        
        found_degrees = set()
        for pattern, degree_type in degree_patterns:
            matches = re.findall(pattern, search_text)
            if matches:
                if degree_type not in found_degrees:
                    found_degrees.add(degree_type)
                    field = ""
                    if isinstance(matches[0], str) and len(matches[0].strip()) > 2:
                        # Clean up the field
                        field = matches[0].strip()
                        # Remove common trailing words
                        field = re.sub(r'\s*(from|at|university|college|institute).*', '', field, flags=re.IGNORECASE)
                        field = field.strip()[:40]
                    education_list.append(Education(
                        degree=degree_type,
                        field=field.title() if field else "",
                        institution=""
                    ))
        
        # If no specific degree found, look for general education keywords
        if not education_list:
            if any(word in text_lower for word in ['ph.d', 'phd', 'doctorate']):
                education_list.append(Education(degree="PhD", field="", institution=""))
            elif any(word in text_lower for word in ["master's", 'mba', 'm.s.', 'ms in', 'ma in']):
                education_list.append(Education(degree="Master's", field="", institution=""))
            elif any(word in text_lower for word in ["bachelor's", 'b.s.', 'bs in', 'b.tech']):
                education_list.append(Education(degree="Bachelor's", field="", institution=""))
        
        return education_list
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract name from first lines of resume."""
        lines = text.strip().split('\n')
        for line in lines[:5]:
            line = line.strip()
            # Skip empty lines and common headers
            if not line or len(line) < 3:
                continue
            if any(word in line.lower() for word in ['resume', 'cv', 'curriculum', 'email', 'phone', '@', 'http']):
                continue
            
            # Handle spaced-out names like "A S I F  K A M A L"
            # Split by double space to get name parts
            if '  ' in line:
                parts = [p.strip() for p in line.split('  ') if p.strip()]
                if len(parts) >= 2:
                    # Join letters within each part
                    name_parts = []
                    for part in parts:
                        if all(c.isupper() or c.isspace() for c in part):
                            name_parts.append(part.replace(' ', ''))
                        else:
                            name_parts.append(part)
                    if name_parts:
                        return ' '.join(name_parts).title()
            
            # Normal name extraction
            if len(line) < 50 and line[0].isupper():
                name = ' '.join(line.split())
                words = name.split()
                if 2 <= len(words) <= 4 and not any(w.lower() in ['engineer', 'developer', 'manager', 'analyst'] for w in words):
                    return name
        return None

    async def parse_resume(self, file_path: str) -> ResumeData:
        """
        Parse a resume file and extract structured data.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            ResumeData object with extracted information
        """
        # Extract raw text
        raw_text = self.extract_text(file_path)
        logger.info(f"Extracted {len(raw_text)} chars from {file_path}")
        
        if not raw_text or len(raw_text) < 50:
            logger.warning(f"Very little text extracted from {file_path}")
            return ResumeData(raw_text=raw_text)
        
        # Use fast regex extraction (skip slow LLM call for speed)
        name = self._extract_name(raw_text)
        email = self._extract_email(raw_text)
        phone = self._extract_phone(raw_text)
        skills = self._basic_skill_extraction(raw_text)
        linkedin = self._extract_linkedin(raw_text)
        experience_years = self._extract_experience_years(raw_text)
        education = self._extract_education(raw_text)
        
        # Only use LLM if critical data is missing (rare case)
        llm_data = None
        if not name or not email or len(skills) < 3:
            logger.info("Critical data missing - using LLM extraction")
            try:
                llm_data = await self.usf_client.extract_resume_info(raw_text[:4000])
                logger.info(f"LLM extraction result: {llm_data}")
            except Exception as e:
                logger.warning(f"LLM extraction failed: {e}")
        
        # Merge LLM data with basic extraction
        if llm_data:
            name = llm_data.get('name') or name
            email = llm_data.get('email') or email
            phone = llm_data.get('phone') or phone
            
            # Merge skills
            llm_skills = llm_data.get('skills', [])
            if isinstance(llm_skills, list):
                skills = list(set(skills + [s.lower() for s in llm_skills if isinstance(s, str)]))
            
            # Get experience from LLM if available
            llm_exp = llm_data.get('total_experience_years', 0)
            if llm_exp and llm_exp > experience_years:
                experience_years = llm_exp
            
            # Get education from LLM if available
            llm_edu = llm_data.get('education', [])
            if llm_edu and isinstance(llm_edu, list) and len(llm_edu) > len(education):
                education = [
                    Education(
                        degree=e.get('degree', ''),
                        field=e.get('field', ''),
                        institution=e.get('institution', '')
                    ) for e in llm_edu if isinstance(e, dict)
                ]
        
        logger.info(f"Final extraction: name={name}, email={email}, skills={len(skills)}, exp={experience_years}")
        
        # Return extraction result
        return ResumeData(
            name=name,
            email=email,
            phone=phone,
            linkedin=linkedin,
            skills=skills,
            education=education,
            experience=[Experience(
                title="",
                company="",
                duration_months=int(experience_years * 12)
            )] if experience_years > 0 else [],
            raw_text=raw_text
        )
    
    async def parse_resume_from_bytes(
        self, 
        content: bytes, 
        filename: str,
        upload_dir: str
    ) -> Tuple[ResumeData, str]:
        """
        Parse a resume from uploaded bytes.
        
        Args:
            content: File content as bytes
            filename: Original filename
            upload_dir: Directory to save the file
            
        Returns:
            Tuple of (ResumeData, saved_file_path)
        """
        # Ensure upload directory exists
        Path(upload_dir).mkdir(parents=True, exist_ok=True)
        
        # Generate safe filename
        safe_filename = re.sub(r'[^\w\-_\.]', '_', filename)
        file_path = os.path.join(upload_dir, safe_filename)
        
        # Handle duplicate filenames
        base, ext = os.path.splitext(file_path)
        counter = 1
        while os.path.exists(file_path):
            file_path = f"{base}_{counter}{ext}"
            counter += 1
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Parse resume
        resume_data = await self.parse_resume(file_path)
        
        return resume_data, file_path


# Global parser instance
resume_parser = ResumeParser()
